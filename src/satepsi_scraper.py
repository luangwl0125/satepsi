"""
SATEPSI Scraper - Sistema de Monitoramento de Instrumentos Psicológicos

Este módulo fornece funcionalidades para extrair, processar e monitorar dados do SATEPSI
(Sistema de Avaliação de Testes Psicológicos) do Conselho Federal de Psicologia.

Funcionalidades principais:
- Extração de dados do site do SATEPSI
- Processamento e formatação dos dados
- Geração de relatórios em diferentes formatos
- Monitoramento de alterações
- Notificações por e-mail

Exemplo de uso:
    >>> from satepsi_scraper import main
    >>> main()  # Executa o scraping e gera relatórios
"""

import json
import smtplib
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Tuple, Optional

import requests
import pandas as pd
from bs4 import BeautifulSoup
from docx import Document
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText

from config import (
    SATEPSI_URL, FIELD_MAP, EMAIL_CONFIG,
    PREVIOUS_DATA_FILE, DATA_DIR,
    get_timestamp_filename
)

# ── FUNÇÕES DE EXTRAÇÃO ────────────────────────────────────────────────────────

def fetch_from_url() -> List[Dict]:
    """
    Extrai dados diretamente do site do SATEPSI.
    
    Returns:
        List[Dict]: Lista de dicionários contendo os dados dos instrumentos
        
    Raises:
        requests.RequestException: Se houver erro na requisição HTTP
        ValueError: Se o HTML retornado estiver em formato inválido
    """
    resp = requests.get(SATEPSI_URL + "lista_teste_completa.cfm", headers={"User-Agent": "Mozilla/5.0"})
    resp.encoding = "latin1"
    soup = BeautifulSoup(resp.text, "html5lib")
    return parse_html(soup)

def fetch_from_file(file_path: str) -> List[Dict]:
    """
    Extrai dados de um arquivo HTML local.
    
    Args:
        file_path (str): Caminho para o arquivo HTML
        
    Returns:
        List[Dict]: Lista de dicionários contendo os dados dos instrumentos
        
    Raises:
        FileNotFoundError: Se o arquivo não existir
        ValueError: Se o HTML estiver em formato inválido
    """
    with open(file_path, encoding="latin1") as f:
        html = f.read()
    soup = BeautifulSoup(html, "html5lib")
    return parse_html(soup)

def parse_html(soup: BeautifulSoup) -> List[Dict]:
    """
    Processa o HTML e extrai os dados dos instrumentos.
    
    Args:
        soup (BeautifulSoup): Objeto BeautifulSoup com o HTML parseado
        
    Returns:
        List[Dict]: Lista de dicionários contendo os dados dos instrumentos
    """
    instruments = []
    
    # Procura por todos os itens do acordeão
    for item in soup.find_all("div", class_="accordion-item"):
        # O título e status estão no primeiro <a>
        header = item.find("a", recursive=False)
        if not header:
            continue
            
        # Extrai o título e status
        title = header.get_text(strip=True)
        status_badge = header.find("span", class_="badge")
        status = status_badge.get_text(strip=True) if status_badge else ""
        
        # Remove o status do título
        title = title.replace(status, "").strip()
        
        # Inicializa o registro
        rec = {v: "" for v in FIELD_MAP.values()}
        rec["Status"] = status
        rec["Nome"] = title
        
        # Procura por todas as linhas de informação
        for row in item.find_all("div", class_="row"):
            label = row.find("div", class_="font-weight-bold")
            value = row.find("div", class_="col-md-8") or row.find("div", class_="col-md-10")
            
            if not label or not value:
                continue
                
            label_text = label.get_text(strip=True).rstrip(":")
            value_text = value.get_text(" ", strip=True)
            
            # Mapeia o campo
            key = FIELD_MAP.get(label_text)
            if key:
                rec[key] = value_text
                
        instruments.append(rec)
    
    return instruments

def parse_table(table: BeautifulSoup) -> List[Dict]:
    """
    Extrai dados do formato de tabela.
    
    Args:
        table (BeautifulSoup): Elemento table do HTML
        
    Returns:
        List[Dict]: Lista de dicionários com os dados extraídos
    """
    instruments = []
    rows = table.find_all("tr")
    
    for i in range(1, len(rows), 2):
        summary = rows[i]
        detail = rows[i+1] if i+1 < len(rows) else None
        
        cols = summary.find_all("td")
        if len(cols) < 4:
            continue
            
        rec = {v: "" for v in FIELD_MAP.values()}
        rec["Status"] = cols[0].get_text(strip=True)
        rec["Autores"] = cols[1].get_text(strip=True)
        rec["Editora"] = cols[2].get_text(strip=True)
        rec["Construto"] = cols[3].get_text(strip=True)
        
        if detail:
            td = detail.find("td", attrs={"colspan": True})
            if td:
                for s in td.find_all("strong"):
                    label = s.get_text(strip=True).rstrip(":")
                    key = FIELD_MAP.get(label)
                    if not key:
                        continue
                    val = ""
                    nxt = s.next_sibling
                    if nxt:
                        val = nxt.strip()
                    rec[key] = val
                    
        instruments.append(rec)
    
    return instruments

def parse_links(soup: BeautifulSoup) -> List[Dict]:
    """
    Extrai dados do formato de links.
    
    Args:
        soup (BeautifulSoup): Objeto BeautifulSoup com o HTML parseado
        
    Returns:
        List[Dict]: Lista de dicionários com os dados extraídos
    """
    instruments = []
    
    for a in soup.find_all("a", string=lambda t: t and t.strip().startswith("+")):
        rec = {v: "" for v in FIELD_MAP.values()}
        rec["Construto"] = a.get_text(strip=True)[1:].strip()
        
        for sib in a.next_siblings:
            if isinstance(sib, BeautifulSoup.Tag) and sib.name == "a" and sib.get_text(strip=True).startswith("+"):
                break
                
            text = ""
            if isinstance(sib, BeautifulSoup.NavigableString):
                text = sib.strip()
            elif isinstance(sib, BeautifulSoup.Tag):
                text = sib.get_text(" ", strip=True)
                
            if ":" not in text:
                continue
                
            label, val = [s.strip() for s in text.split(":", 1)]
            key = FIELD_MAP.get(label)
            if key:
                rec[key] = val
                
        instruments.append(rec)
    
    return instruments

# ── FUNÇÕES DE PROCESSAMENTO ───────────────────────────────────────────────────

def save_json(data: List[Dict], path: Path) -> None:
    """
    Salva os dados em formato JSON.
    
    Args:
        data (List[Dict]): Dados a serem salvos
        path (Path): Caminho do arquivo de saída
    """
    with open(path, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)

def load_json(path: Path) -> List[Dict]:
    """
    Carrega dados de um arquivo JSON.
    
    Args:
        path (Path): Caminho do arquivo JSON
        
    Returns:
        List[Dict]: Dados carregados
        
    Raises:
        FileNotFoundError: Se o arquivo não existir
        json.JSONDecodeError: Se o JSON estiver mal formatado
    """
    with open(path, encoding="utf-8") as f:
        return json.load(f)

def save_excel(data: List[Dict], path: Path) -> None:
    """
    Salva os dados em formato Excel.
    
    Args:
        data (List[Dict]): Dados a serem salvos
        path (Path): Caminho do arquivo de saída
    """
    # Garante que a coluna 'Nome' seja a primeira
    cols = ["Nome"] + [c for c in FIELD_MAP.values() if c != "Nome"]
    df = pd.DataFrame(data, columns=cols)
    df.to_excel(path, index=False)

def generate_word(data: List[Dict], path: Path) -> None:
    """
    Gera um relatório em formato Word.
    
    Args:
        data (List[Dict]): Dados a serem incluídos no relatório
        path (Path): Caminho do arquivo de saída
    """
    doc = Document()
    doc.add_heading(f"Instrumentos SATEPSI ({datetime.today():%d/%m/%Y})", level=1)
    
    for inst in data:
        doc.add_heading(inst["Construto"], level=2)
        for k, v in inst.items():
            if v:  # Só adiciona campos não vazios
                doc.add_paragraph(f"{k}: {v}")
    
    doc.save(path)

# ── FUNÇÕES DE MONITORAMENTO ───────────────────────────────────────────────────

def diff_status(old: List[Dict], new: List[Dict]) -> List[Tuple[str, str, str]]:
    """
    Compara o status dos instrumentos entre duas versões.
    
    Args:
        old (List[Dict]): Versão anterior dos dados
        new (List[Dict]): Versão atual dos dados
        
    Returns:
        List[Tuple[str, str, str]]: Lista de tuplas (nome, status_antigo, status_novo)
    """
    old_map = {i["Construto"]: i.get("Status", "") for i in old}
    return [
        (i["Construto"], old_map[i["Construto"]], i.get("Status", ""))
        for i in new
        if i["Construto"] in old_map and old_map[i["Construto"]] != i.get("Status", "")
    ]

def send_email(changes: List[Tuple[str, str, str]]) -> None:
    """
    Envia e-mail com as alterações detectadas.
    
    Args:
        changes (List[Tuple[str, str, str]]): Lista de alterações detectadas
        
    Note:
        Só envia e-mail se houver alterações e se as configurações de e-mail
        estiverem disponíveis
    """
    if not EMAIL_CONFIG:
        print("[AVISO] Configurações de e-mail faltando: SMTP_HOST, SMTP_USER, SMTP_PASS, EMAIL_FROM, EMAIL_TO")
        return
        
    try:
        # Configura servidor SMTP
        server = smtplib.SMTP(EMAIL_CONFIG["SMTP_HOST"], EMAIL_CONFIG["SMTP_PORT"])
        server.starttls()
        server.login(EMAIL_CONFIG["SMTP_USER"], EMAIL_CONFIG["SMTP_PASS"])
        
        # Prepara mensagem
        msg = MIMEMultipart()
        msg["From"] = EMAIL_CONFIG["EMAIL_FROM"]
        msg["To"] = EMAIL_CONFIG["EMAIL_TO"]
        msg["Subject"] = "Atualização Satepsi"
        
        # Corpo do e-mail
        body = f"""Olá Luan,

"""
        if changes:
            body += "Foram detectadas as seguintes alterações:\n\n"
            for name, old, new in changes:
                body += f"- {name}: {old} -> {new}\n"
        else:
            body += "Nenhuma alteração foi detectada.\n"
            
        body += """
Atenciosamente,
Sistema de Monitoramento Satepsi"""

        # Adiciona corpo com codificação UTF-8
        msg.attach(MIMEText(body.encode('utf-8'), 'plain', 'utf-8'))
        
        # Envia e-mail
        server.send_message(msg)
        server.quit()
        print("[SUCESSO] E-mail enviado com sucesso!")
        
    except Exception as e:
        print(f"[ERRO] Erro ao enviar e-mail: {str(e)}")

# ── FUNÇÃO PRINCIPAL ───────────────────────────────────────────────────────────

def main():
    """
    Função principal do script.
    
    Fluxo de execução:
    1. Extrai dados do site do SATEPSI
    2. Salva os dados em diferentes formatos (JSON, Excel, Word)
    3. Compara com a versão anterior
    4. Envia notificações se houver alterações
    5. Atualiza o histórico
    
    Raises:
        Exception: Se houver erro na extração dos dados
    """
    # 1) Extrai dados
    try:
        instruments = fetch_from_url()
    except Exception as e:
        print(f"Erro ao extrair dados do site: {e}")
        return
        
    if not instruments:
        print("[AVISO] Nenhum instrumento encontrado.")
        return
        
    print(f"[SUCESSO] {len(instruments)} instrumentos extraídos com sucesso.")
    
    # 2) Salva dados em diferentes formatos
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    
    # JSON
    json_path = get_timestamp_filename("satepsi_data", "json")
    save_json(instruments, json_path)
    print(f"[SUCESSO] Dados salvos em JSON: {json_path}")
    
    # Excel
    excel_path = get_timestamp_filename("satepsi_data", "xlsx")
    save_excel(instruments, excel_path)
    print(f"[SUCESSO] Dados salvos em Excel: {excel_path}")
    
    # Word
    word_path = get_timestamp_filename("satepsi_report", "docx")
    generate_word(instruments, word_path)
    print(f"[SUCESSO] Relatório gerado em Word: {word_path}")
    
    # 3) Verifica alterações
    if PREVIOUS_DATA_FILE.exists():
        old_data = load_json(PREVIOUS_DATA_FILE)
        changes = diff_status(old_data, instruments)
        
        if changes:
            print("\nAlterações detectadas:")
            for name, old, new in changes:
                print(f"- {name}: {old} -> {new}")
        
        # Envia e-mail independente de haver alterações
        send_email(changes)
    else:
        # Primeira execução, envia e-mail informando que não há alterações
        send_email([])
    
    # 4) Atualiza histórico
    save_json(instruments, PREVIOUS_DATA_FILE)
    print(f"[SUCESSO] Histórico atualizado: {PREVIOUS_DATA_FILE}")

if __name__ == "__main__":
    main() 
